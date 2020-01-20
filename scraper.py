import requests
import urllib.request
import time
from bs4 import BeautifulSoup
from docx import Document

#Opens .docx file and initialises list of articles which cannot be added
document = Document()
failedArticles = []

#Performs main function of parsing HTML, 1 run per day
def mainParser():
    url = 'https://www.abc.net.au/news/archive/?date=' + str(year) + '-' + str("{:02d}".format(month)) + '-' + str("{:02d}".format(day))
    response = requests.get(url)

    soup = BeautifulSoup(response.text, "html.parser")

    pageLinks = len(soup.find_all("li", class_="pagenum"))
    counter = 1
    
    #Performing for each page per date
    while counter <= pageLinks:
        url = 'https://www.abc.net.au/news/archive/?date=' + str(year) + '-' + str("{:02d}".format(month)) + '-' + str("{:02d}".format(day)) + '&page=' + str(counter)
        response = requests.get(url)

        soup = BeautifulSoup(response.text, "html.parser")

        potentialArticles = soup.find_all("div", class_="view-textlink")
        
        #Performing for every article on a page
        for potentialArticle in potentialArticles:
            
            #Most articles have all following attributes: title, timestamp and content
            try:
                if potentialArticle.parent.name == 'h3':
                    link = potentialArticle.find("a")
                    articleURL = 'https://www.abc.net.au' + link.get('href')
                    articleResponse = requests.get(articleURL)
                    
                    #Get the headline of the article
                    articleSoup = BeautifulSoup(articleResponse.text, "html.parser")
                    articleTitle = articleSoup.find('h1')
                    document.add_heading(articleTitle.getText(), level=1)

                    #Get the timestamp of the article
                    articleContent = articleSoup.find('div', class_='article section')
                    articleDate = articleContent.find('p', class_='published').find('span', class_='timestamp')
                    document.add_paragraph(articleDate.getText())

                    #Get all content for the article, paragraph by paragraph
                    articleText = articleContent.find_all('p', class_=False)
                    for p in articleText:
                        document.add_paragraph(p.getText()) 

                    document.add_paragraph()
            
            #Articles which dont have one of the above attributes have their titles appended to the failedArticles list
            except:
                failedArticles.append(potentialArticle.getText())

        counter += 1

#Getting all articles between 1/1/2009 & 31/12/2019
year = 2009
targetYear = 2019
month = 1
day = 1

monthsOf31 = [1, 3, 5, 7, 8, 10, 12]
monthsOf30 = [4, 6, 9, 11]

while year <= targetYear:
    while month <= 12:
        if month in monthsOf31:
            daysInMonth = 31
        elif month in monthsOf30:
            daysInMonth = 30
        elif month == 2:
            if month % 4 == 0:
                daysInMonth = 29
            else:
                daysInMonth = 28

        while day <= daysInMonth:
            print(str(day) + '/' + str(month) + '/' + str(year))
            mainParser()
            day += 1

        month += 1
        day = 1
    year += 1
    month = 1

#Adding failed articles and saving the document
failedArticlesParagraph = document.add_paragraph()
for failedArticle in failedArticles:
    newLine = document.add_run(failedArticle)
    newLine.add_break()
document.save('results.docx')
    